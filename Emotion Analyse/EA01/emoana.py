from collections import defaultdict
import os
import re
import jieba
import codecs

import pandas as pd
from docx import Document

# 生成stopword表，需要去除一些否定词和程度词汇
stopwords = set()
fr = open('stopwords.txt', 'r', encoding='utf-8')
for word in fr:
    stopwords.add(word.strip())  # Python strip() 方法用于移除字符串头尾指定的字符（默认为空格或换行符）或字符序列。
# 读取否定词文件
not_word_file = open('noword.txt', 'r+', encoding='utf-8')
not_word_list = not_word_file.readlines()
not_word_list = [w.strip() for w in not_word_list]
# 读取程度副词文件
with open("adverb.txt", "r", encoding="utf-8") as degree_file:
    degree_list = degree_file.readlines()  # 读取所有行
    degree_list = [item.strip().split(',')[0] for item in degree_list]  # 处理数据
# 生成新的停用词表
with open('stopwords.txt', 'w', encoding='utf-8') as f:
    for word in stopwords:
        if (word not in not_word_list) and (word not in degree_list):
            f.write(word + '\n')


# jieba分词后去除停用词
def seg_word(sentence):
    seg_list = jieba.cut(sentence)
    seg_result = []
    for i in seg_list:
        seg_result.append(i)

    stopwords = set()
    with open('stopwords.txt', 'r', encoding='utf-8') as fr:  # 关键修改：指定 encoding='utf-8'
        for line in fr:
            stopwords.add(line.strip())

    return list(filter(lambda x: x not in stopwords, seg_result))


# 找出文本中的情感词、否定词和程度副词
from collections import defaultdict  # 确保在文件顶部导入


def classify_words(word_list):
    # 读取情感词典文件
    with open('sentiment_score.txt', 'r', encoding='utf-8') as sen_file:
        sen_list = sen_file.readlines()

    # 创建情感字典（使用普通字典即可，不需要defaultdict）
    sen_dict = {}
    for line in sen_list:
        parts = line.strip().split(' ')
        if len(parts) == 2:
            sen_dict[parts[0]] = float(parts[1])  # 将权重转换为float

    # 读取否定词文件
    with open('noword.txt', 'r', encoding='utf-8') as not_word_file:
        not_word_list = [line.strip() for line in not_word_file]

    # 读取程度副词文件
    with open('adverb.txt', 'r', encoding='utf-8') as degree_file:
        degree_dict = {}
        for line in degree_file:
            parts = line.strip().split(',')
            if len(parts) >= 2:
                degree_dict[parts[0]] = float(parts[1])

    # 分类词语
    sen_words = {}
    not_words = {}
    degree_words = {}

    for idx, word in enumerate(word_list):
        if word in sen_dict and word not in not_word_list and word not in degree_dict:
            sen_words[idx] = sen_dict[word]
        elif word in not_word_list and word not in degree_dict:
            not_words[idx] = -1
        elif word in degree_dict:
            degree_words[idx] = degree_dict[word]

    return sen_words, not_words, degree_words

    # 读取否定词文件
    from collections import defaultdict

    def classify_words(word_list, sen_dict):
        # 读取否定词文件
        with open('noword.txt', 'r', encoding='utf-8') as not_word_file:
            not_word_list = [line.strip() for line in not_word_file]

        # 读取程度副词文件（确保使用utf-8编码）
        with open('adverb.txt', 'r', encoding='utf-8') as degree_file:
            degree_list = [line.strip() for line in degree_file]

        # 构建程度副词字典
        degree_dict = {}
        for item in degree_list:
            parts = item.split(',')
            if len(parts) >= 2:  # 确保有足够的部分
                degree_dict[parts[0]] = float(parts[1])

        # 初始化分类字典
        sen_word = {}
        not_word = {}
        degree_word = {}

        # 分类处理
        for i, word in enumerate(word_list):
            if word in sen_dict and word not in not_word_list and word not in degree_dict:
                sen_word[i] = sen_dict[word]
            elif word in not_word_list and word not in degree_dict:
                not_word[i] = -1
            elif word in degree_dict:
                degree_word[i] = degree_dict[word]

        return sen_word, not_word, degree_word

    # 关闭打开的文件
    sen_file.close()
    not_word_file.close()
    degree_file.close()
    # 返回分类结果
    return sen_word, not_word, degree_word


# 计算情感词的分数
def score_sentiment(sen_word, not_word, degree_word, seg_result):
    # 权重初始化为1
    W = 1
    score = 0
    # 情感词下标初始化
    sentiment_index = -1
    # 情感词的位置下标集合
    sentiment_index_list = list(sen_word.keys())
    # 遍历分词结果
    for i in range(0, len(seg_result)):
        # 如果是情感词
        if i in sen_word.keys():
            # 权重*情感词得分
            score += W * float(sen_word[i])
            # 情感词下标加一，获取下一个情感词的位置
            sentiment_index += 1
            if sentiment_index < len(sentiment_index_list) - 1:
                # 判断当前的情感词与下一个情感词之间是否有程度副词或否定词
                for j in range(sentiment_index_list[sentiment_index], sentiment_index_list[sentiment_index + 1]):
                    # 更新权重，如果有否定词，权重取反
                    if j in not_word.keys():
                        W *= -1
                    elif j in degree_word.keys():
                        W *= float(degree_word[j])
        # 定位到下一个情感词
        if sentiment_index < len(sentiment_index_list) - 1:
            i = sentiment_index_list[sentiment_index + 1]
    return score


# 计算得分
def sentiment_score(sentence):
    # 1.对文档分词
    seg_list = seg_word(sentence)
    # 2.将分词结果转换成字典，找出情感词、否定词和程度副词
    sen_word, not_word, degree_word = classify_words(seg_list)
    # 3.计算得分
    score = score_sentiment(sen_word, not_word, degree_word, seg_list)
    return score


print("我今天很高兴也非常开心    ", sentiment_score("我今天很高兴也非常开心"))
print('天灰蒙蒙的，路上有只流浪狗，旁边是破旧不堪的老房子   ',
      sentiment_score('天灰蒙蒙的，路上有只流浪狗，旁边是破旧不堪的老房子'))
print('愤怒、悲伤和埋怨解决不了问题    ', sentiment_score('愤怒、悲伤和埋怨解决不了问题'))
print('要每天都开心快乐    ', sentiment_score('要每天都开心快乐'))
print('我不喜欢这个世界，我只喜欢你    ', sentiment_score('我不喜欢这个世界，我只喜欢你'))


from docx import Document


def analyze_document_sentiment(file_path):
    # 打开Word文档
    doc = Document(file_path)

    # 遍历文档中的每个段落
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # 只处理非空段落
            print(f"段落 {i + 1}: {text}    ", sentiment_score(text))


def analyze_document_and_export_to_excel(file_path, output_excel_path):
    """
    分析Word文档中的每个段落，并将段落文本及其情感得分导出到Excel文件。

    Args:
        file_path (str): 输入的Word文档路径。
        output_excel_path (str): 输出的Excel文件路径。
    """
    try:
        doc = Document(file_path)
    except FileNotFoundError:
        print(f"错误：找不到文件 '{file_path}'。请检查文件路径是否正确。")
        return

    # 创建一个空的列表来存储数据
    data = []

    # 遍历文档中的每个段落
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # 只处理非空段落
            score = sentiment_score(text)
            # 将段落和得分作为一行数据添加到列表中
            data.append({
                '段落序号': i + 1,
                '段落文本': text,
                '情感得分': score
            })

    # 使用pandas将数据列表转换为DataFrame
    df = pd.DataFrame(data)

    # 将DataFrame写入Excel文件
    try:
        df.to_excel(output_excel_path, index=False)
        print(f"成功将情感分析结果导出到 '{output_excel_path}'。")
    except Exception as e:
        print(f"导出Excel时发生错误：{e}")

# ----------------- 使用示例 -----------------
# 定义输入和输出文件路径
file_path = "../荆楚名校笔记.docx"
output_excel_path = "情感分析结果.xlsx"

# 调用新的函数来执行分析并导出
analyze_document_and_export_to_excel(file_path, output_excel_path)
